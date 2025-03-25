#!/usr/bin/env python3
"""
Web to PDF Converter Tool
-------------------------
This script crawls a website recursively and converts its content into PDF, HTML, JSON, or DOCX format.
It provides various options for controlling the crawling behavior, filtering content, and formatting output.

Features:
- Automatic link discovery and following
- Multiple output formats (PDF, HTML, JSON, DOCX)
- Depth control for limiting crawl scope
- Content filtering by keywords
- Category filtering
- SVG image handling
- Interactive mode for URL selection
- Sitemap-based crawling
- Table of contents generation (PDF)

Usage Examples:
--------------
# Basic usage with default settings
python web_to_doc.py --url https://docs.example.com/

# Specify output file
python web_to_doc.py --url https://docs.example.com/ --output documentation.pdf

# Limit crawling to 2 levels deep
python web_to_doc.py --url https://docs.example.com/ --max-depth 2

# Convert to HTML format
python web_to_doc.py --url https://docs.example.com/ --format html

# Convert to Markdown format
python web_to_doc.py --url https://docs.example.com/ --format md

# Only include pages containing certain keywords
python web_to_doc.py --url https://docs.example.com/ --contains "guide,tutorial,reference"

# Interactive mode to select which URLs to process
python web_to_doc.py --url https://docs.example.com/ --interactive

# Generate table of contents
python web_to_doc.py --url https://docs.example.com/ --toc

Dependencies:
- requests: For HTTP requests
- beautifulsoup4: For HTML parsing
- reportlab: For PDF generation
- Pillow: For image processing
- (optional) python-docx: For DOCX output
- (optional) lxml: For XML parsing

Author: Created with assistance from Claude AI
Date: March 2025
"""

import requests
from bs4 import BeautifulSoup
import os
import io
import tempfile
import re
import logging
import time
import argparse
import json
import sys
from urllib.parse import urljoin, urlparse
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Preformatted, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.units import inch
from PIL import Image as PILImage

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class WebToPdfConverter:
    def __init__(self, base_url, output_path, options=None):
        self.base_url = base_url
        self.output_path = output_path
        self.domain = urlparse(base_url).netloc
        self.base_path = urlparse(base_url).path.rsplit('/', 1)[0] + '/'
        self.temp_dir = tempfile.mkdtemp()
        self.image_counter = 0
        self.pdf_elements = []
        self.visited_urls = set()
        self.to_visit = [base_url]
        self.contents = {}  # Store page content for non-PDF formats
        
        # Default options
        self.options = {
            'max_pages': 250,
            'max_depth': None,
            'delay': 1,
            'timeout': 10,
            'format': 'pdf',
            'use_sitemap': False,
            'sitemap_url': None,
            'categories': None,
            'contains': None,
            'not_contains': None,
            'create_toc': False,
            'interactive': False
        }
        
        # Update with custom options
        if options:
            self.options.update(options)
        
        # URL depth for max_depth tracking
        self.url_depth = {base_url: 0}
        
        # Sitemap URLs, if used
        self.sitemap_urls = []
    
    def download_page(self, url):
        """Downloads a webpage and returns the HTML content"""
        try:
            logger.info(f"Loading page: {url}")
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, headers=headers, timeout=self.options['timeout'])
            response.raise_for_status()
            
            # Check the Content-Type of the response
            content_type = response.headers.get('Content-Type', '').lower()
            if 'text/html' not in content_type and 'application/xhtml+xml' not in content_type:
                logger.warning(f"Skipping non-HTML content: {url} (Content-Type: {content_type})")
                return None
                
            return response.text
        except requests.exceptions.RequestException as e:
            logger.error(f"Error loading page {url}: {e}")
            return None    def is_valid_url(self, url, from_url=None):
        """Checks if a URL is valid and belongs to the same documentation area"""
        parsed = urlparse(url)
        
        # Only URLs to the same domain
        if parsed.netloc and parsed.netloc != self.domain:
            return False
            
        # Only URLs in the same documentation area
        path = parsed.path
        if not path.startswith(self.base_path):
            return False
            
        # No anchor links or query parameters
        if "#" in url:
            url = url.split("#")[0]
            
        # No files that aren't HTML (like PDFs, images, JSON, etc.)
        if path.endswith(('.pdf', '.png', '.jpg', '.jpeg', '.gif', '.svg', '.zip', '.exe', '.json', '.xml', '.js', '.css')):
            return False
        
        # Check depth limitation, if set
        if self.options['max_depth'] is not None and from_url:
            parent_depth = self.url_depth.get(from_url, 0)
            if parent_depth >= self.options['max_depth']:
                return False
            
            # Set depth for new URL
            self.url_depth[url] = parent_depth + 1
            
        return True
    
    def extract_links(self, soup, current_url):
        """Extracts all links from a page"""
        links = []
        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href']
            absolute_url = urljoin(current_url, href)
            
            if self.is_valid_url(absolute_url, current_url) and absolute_url not in self.visited_urls:
                links.append(absolute_url)
                
        return links
    
    def download_image(self, img_url, base_url):
        """Downloads an image and saves it temporarily"""
        try:
            absolute_url = urljoin(base_url, img_url)
            response = requests.get(absolute_url, stream=True, timeout=self.options['timeout'])
            response.raise_for_status()
            
            img_data = response.content
            
            # Check if it's an SVG image
            is_svg = img_url.lower().endswith('.svg') or 'image/svg+xml' in response.headers.get('Content-Type', '').lower()
            
            if is_svg:
                # For SVG files, create a simple PNG placeholder
                # Alternatively, install cairosvg with: pip install cairosvg
                img_path = os.path.join(self.temp_dir, f"img_{self.image_counter}.png")
                self.image_counter += 1
                
                # Create a blank image for SVG placeholder
                placeholder = PILImage.new('RGB', (300, 100), color=(240, 240, 240))
                from PIL import ImageDraw
                draw = ImageDraw.Draw(placeholder)
                draw.text((10, 40), f"SVG Image: {os.path.basename(img_url)}", fill=(0, 0, 0))
                placeholder.save(img_path, "PNG")
            else:
                # Process normal image
                img_path = os.path.join(self.temp_dir, f"img_{self.image_counter}.png")
                self.image_counter += 1
                
                # Convert and save the image
                img = PILImage.open(io.BytesIO(img_data))
                img.save(img_path, "PNG")
            
            return img_path
        except Exception as e:
            logger.error(f"Error downloading image {img_url}: {e}")
            return None
    
    def check_content_filters(self, soup, url):
        """Checks if the content meets the filter criteria"""
        # If no filters are set, accept everything
        if not self.options['contains'] and not self.options['not_contains']:
            return True
            
        # Get all text from the page
        page_text = soup.get_text().lower()
        
        # Check "contains" filter
        if self.options['contains']:
            contains_keywords = [k.lower().strip() for k in self.options['contains'].split(',')]
            if not any(keyword in page_text for keyword in contains_keywords):
                logger.info(f"Page doesn't contain any of the required keywords: {url}")
                return False
                
        # Check "not_contains" filter
        if self.options['not_contains']:
            not_contains_keywords = [k.lower().strip() for k in self.options['not_contains'].split(',')]
            if any(keyword in page_text for keyword in not_contains_keywords):
                logger.info(f"Page contains excluded keywords: {url}")
                return False
                
        return True
    
    def get_category(self, soup):
        """Attempts to determine the category of a page"""
        # Depending on the page structure, different selectors can be used
        category_selectors = [
            '.breadcrumbs', 
            '.nav-item.active',
            '.sidebar .active',
            'header .category'
        ]
        
        for selector in category_selectors:
            category_elem = soup.select_one(selector)
            if category_elem:
                return category_elem.text.strip()
                
        return None
    
    def parse_sitemap(self, sitemap_url=None):
        """Parses the sitemap and extracts URLs"""
        if not sitemap_url:
            # Try to find standard sitemap paths
            sitemap_candidates = [
                f"{self.base_url}/sitemap.xml",
                f"https://{self.domain}/sitemap.xml",
                f"https://{self.domain}/sitemap_index.xml"
            ]
            
            for candidate in sitemap_candidates:
                try:
                    response = requests.get(candidate, timeout=self.options['timeout'])
                    if response.status_code == 200:
                        sitemap_url = candidate
                        break
                except:
                    continue
                    
        if not sitemap_url:
            logger.error("No sitemap found.")
            return []
            
        try:
            logger.info(f"Loading sitemap: {sitemap_url}")
            response = requests.get(sitemap_url, timeout=self.options['timeout'])
            response.raise_for_status()
            
            # Try to use lxml-xml parser, fall back to html.parser if not available
            try:
                soup = BeautifulSoup(response.text, 'lxml-xml')
            except:
                soup = BeautifulSoup(response.text, 'html.parser')
                
            urls = []
            
            # Process sitemap index (contains links to other sitemaps)
            sitemapindex = soup.find('sitemapindex')
            if sitemapindex:
                for sitemap in sitemapindex.find_all('sitemap'):
                    loc = sitemap.find('loc')
                    if loc:
                        urls.extend(self.parse_sitemap(loc.text))
            else:
                # Process normal sitemap
                for url in soup.find_all('url'):
                    loc = url.find('loc')
                    if loc and self.is_valid_url(loc.text):
                        urls.append(loc.text)
                        
            return urls
            
        except Exception as e:
            logger.error(f"Error parsing sitemap: {e}")
            return []
            
    def process_page(self, url):
        """Processes a single page and extracts content for the PDF"""
        html_content = self.download_page(url)
        if not html_content:
            return []
        
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Check content filter
            if not self.check_content_filters(soup, url):
                return []
                
            # Check category filter
            if self.options['categories']:
                page_category = self.get_category(soup)
                if page_category:
                    categories = [c.strip() for c in self.options['categories'].split(',')]
                    if not any(cat.lower() in page_category.lower() for cat in categories):
                        logger.info(f"Page doesn't belong to the selected categories: {url}")
                        return []
            
            # Remove navigation, footer, etc.
            for selector in ['.navbar', '.footer', '.sidebar', 'nav', '.cookie-banner', '.announcement', '.header']:
                for element in soup.select(selector):
                    element.decompose()
            
            # Extract the main content
            main_content = soup.select_one('main') or soup.select_one('article') or soup.select_one('.content')
            if not main_content:
                main_content = soup.body
                
            if not main_content:
                logger.warning(f"No main content found in: {url}")
                return []
            
            # Store page content for non-PDF formats
            self.contents[url] = {
                'title': soup.title.string if soup.title else url.split('/')[-1],
                'html': str(main_content),
                'text': main_content.get_text(),
                'url': url
            }
            
            # Add page title
            title = soup.title.string if soup.title else url.split('/')[-1]
            elements = [
                Paragraph(f"<b>{title}</b>", ParagraphStyle('Title', fontSize=16, spaceAfter=12)),
                Paragraph(f"Source: {url}", ParagraphStyle('URL', fontSize=8, textColor=colors.gray, spaceAfter=12))
            ]
            
            # Styles for various elements
            styles = getSampleStyleSheet()
            header_style = ParagraphStyle('Header', fontSize=14, spaceAfter=8, spaceBefore=12)
            text_style = ParagraphStyle('Text', fontSize=10, spaceAfter=8, leading=14)
            code_style = ParagraphStyle('Code', fontName='Courier', fontSize=8, spaceAfter=8, 
                                       backColor=colors.lightgrey, borderWidth=1, borderColor=colors.lightgrey,
                                       borderPadding=5, leading=12)
            
            # Extract links for further visits, if not in sitemap mode
            if not self.options['use_sitemap']:
                new_links = self.extract_links(soup, url)
                for link in new_links:
                    if link not in self.visited_urls and link not in self.to_visit:
                        self.to_visit.append(link)            # Process elements of the main content
            for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'code', 'img', 'div']):
                # Headings
                if element.name in ['h1', 'h2', 'h3', 'h4']:
                    text = element.text.strip()
                    if text:
                        elements.append(Paragraph(text, header_style))
                
                # Paragraphs
                elif element.name == 'p':
                    text = element.text.strip()
                    if text:
                        elements.append(Paragraph(text, text_style))
                
                # Code blocks
                elif element.name == 'pre' or (element.name == 'div' and 'code' in element.get('class', [])):
                    code_text = element.text.strip()
                    if code_text:
                        elements.append(Preformatted(code_text, code_style))
                
                # Inline code
                elif element.name == 'code' and element.parent.name != 'pre':
                    code_text = element.text.strip()
                    if code_text:
                        elements.append(Preformatted(code_text, code_style))
                
                # Images
                elif element.name == 'img' and element.get('src'):
                    img_path = self.download_image(element['src'], url)
                    if img_path:
                        try:
                            img = Image(img_path)
                            # Limit image size
                            max_width = 450
                            if img.drawWidth > max_width:
                                ratio = max_width / img.drawWidth
                                img.drawWidth = max_width
                                img.drawHeight *= ratio
                            elements.append(img)
                            elements.append(Spacer(1, 6))
                        except Exception as e:
                            logger.error(f"Error adding image to PDF: {e}")
            
            # Add a page break at the end
            elements.append(PageBreak())
            
            return elements
            
        except Exception as e:
            logger.error(f"Error processing page {url}: {e}")
            return []
    
    def interactive_mode(self):
        """Allows user to interactively select which URLs to process"""
        if self.options['use_sitemap']:
            # Get all URLs from sitemap
            urls = self.parse_sitemap(self.options['sitemap_url'])
        else:
            # First, crawl the site to discover URLs without processing them
            temp_visited = set()
            to_visit = [self.base_url]
            discovered_urls = []
            
            while to_visit and len(temp_visited) < self.options['max_pages']:
                current_url = to_visit.pop(0)
                if current_url in temp_visited:
                    continue
                
                temp_visited.add(current_url)
                discovered_urls.append(current_url)
                
                html_content = self.download_page(current_url)
                if html_content:
                    soup = BeautifulSoup(html_content, 'html.parser')
                    new_links = self.extract_links(soup, current_url)
                    
                    for link in new_links:
                        if link not in temp_visited and link not in to_visit:
                            to_visit.append(link)
                
                time.sleep(self.options['delay'])
            
            urls = discovered_urls
        
        # Let the user select URLs
        print("\nDiscovered URLs:")
        for i, url in enumerate(urls, 1):
            print(f"{i}. {url}")
        
        print("\nEnter the numbers of the URLs you want to include (comma-separated),")
        print("or 'all' to include all, or 'q' to quit:")
        
        selection = input("> ")
        
        if selection.lower() == 'q':
            sys.exit(0)
        
        if selection.lower() == 'all':
            selected_urls = urls
        else:
            try:
                indices = [int(i.strip()) - 1 for i in selection.split(',')]
                selected_urls = [urls[i] for i in indices if 0 <= i < len(urls)]
            except:
                logger.error("Invalid selection. Exiting.")
                sys.exit(1)
        
        return selected_urls
    
    def create_pdf(self):
        """Creates a PDF from the collected elements"""
        if not self.pdf_elements:
            logger.error("No content to create PDF.")
            return False
        
        try:
            doc = SimpleDocTemplate(self.output_path, pagesize=A4)
            
            # Add table of contents if requested
            if self.options['create_toc']:
                toc = TableOfContents()
                toc.levelStyles = [
                    ParagraphStyle(name='TOCHeading1', fontSize=14, leading=16),
                    ParagraphStyle(name='TOCHeading2', fontSize=12, leading=14, leftIndent=20),
                    ParagraphStyle(name='TOCHeading3', fontSize=10, leading=12, leftIndent=40),
                    ParagraphStyle(name='TOCHeading4', fontSize=10, leading=12, leftIndent=60)
                ]
                
                # Add TOC at the beginning
                self.pdf_elements.insert(0, toc)
                self.pdf_elements.insert(1, PageBreak())
            
            doc.build(self.pdf_elements)
            logger.info(f"PDF created: {self.output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating PDF: {e}")
            return False
    
    def create_html(self):
        """Creates an HTML file from the collected content"""
        try:
            with open(self.output_path, 'w', encoding='utf-8') as f:
                f.write('<!DOCTYPE html>\n<html>\n<head>\n')
                f.write('<meta charset="UTF-8">\n')
                f.write(f'<title>Documentation: {self.base_url}</title>\n')
                f.write('<style>\n')
                f.write('body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }\n')
                f.write('h1 { color: #333; }\n')
                f.write('pre { background-color: #f5f5f5; padding: 10px; border-radius: 5px; }\n')
                f.write('.source { color: #666; font-size: 0.8em; margin-bottom: 20px; }\n')
                f.write('</style>\n')
                f.write('</head>\n<body>\n')
                
                # Add table of contents
                if self.options['create_toc'] and self.contents:
                    f.write('<h2>Table of Contents</h2>\n<ul>\n')
                    for url, content in self.contents.items():
                        f.write(f'<li><a href="#{url.split("//")[1].replace("/", "_")}">{content["title"]}</a></li>\n')
                    f.write('</ul>\n<hr>\n')
                
                # Add content of each page
                for url, content in self.contents.items():
                    f.write(f'<div id="{url.split("//")[1].replace("/", "_")}">\n')
                    f.write(f'<h1>{content["title"]}</h1>\n')
                    f.write(f'<div class="source">Source: <a href="{url}">{url}</a></div>\n')
                    f.write(content['html'])
                    f.write('\n<hr>\n</div>\n')
                
                f.write('</body>\n</html>')
                
            logger.info(f"HTML created: {self.output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating HTML: {e}")
            return False
    
    def create_json(self):
        """Creates a JSON file from the collected content"""
        try:
            output = {
                'base_url': self.base_url,
                'pages': []
            }
            
            for url, content in self.contents.items():
                output['pages'].append({
                    'url': url,
                    'title': content['title'],
                    'text': content['text']
                })
                
            with open(self.output_path, 'w', encoding='utf-8') as f:
                json.dump(output, f, indent=2)
                
            logger.info(f"JSON created: {self.output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating JSON: {e}")
            return False
    
    def create_markdown(self):
        """Creates a Markdown file from the collected content"""
        try:
            with open(self.output_path, 'w', encoding='utf-8') as f:
                # Write title
                f.write(f"# Documentation: {self.base_url}\n\n")
                
                # Add table of contents if requested
                if self.options['create_toc'] and self.contents:
                    f.write("## Table of Contents\n\n")
                    for url, content in self.contents.items():
                        # Create a clean anchor name from the title
                        anchor = content['title'].lower().replace(' ', '-').replace(':', '').replace('.', '')
                        f.write(f"- [{content['title']}](#{anchor})\n")
                    f.write("\n---\n\n")
                
                # Add content of each page
                for url, content in self.contents.items():
                    # Create a clean anchor name from the title
                    anchor = content['title'].lower().replace(' ', '-').replace(':', '').replace('.', '')
                    f.write(f"## {content['title']} <a id=\"{anchor}\"></a>\n\n")
                    f.write(f"Source: [{url}]({url})\n\n")
                    
                    # Convert HTML content to Markdown-friendly text
                    # This is a simple approach, consider using html2text for better conversion
                    text = content['text']
                    
                    # Process text to make it more markdown-friendly
                    # Split into paragraphs and add line breaks
                    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                    for p in paragraphs:
                        f.write(f"{p}\n\n")
                    
                    f.write("---\n\n")
                
            logger.info(f"Markdown created: {self.output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating Markdown: {e}")
            return False
    
    def create_docx(self):
        """Creates a DOCX file from the collected content"""
        try:
            # Import docx only when needed
            from docx import Document
            from docx.shared import Inches, Pt
            
            doc = Document()
            
            # Title
            doc.add_heading(f'Documentation: {self.base_url}', 0)
            
            # Table of contents (just a heading, as Word generates TOC automatically)
            if self.options['create_toc']:
                doc.add_heading('Table of Contents', 1)
                doc.add_paragraph('To update this table, right-click and select "Update Field".')
                doc.add_paragraph()
            
            # Content of each page
            for url, content in self.contents.items():
                doc.add_heading(content['title'], 1)
                doc.add_paragraph(f"Source: {url}").italic = True
                
                # Add text content
                doc.add_paragraph(content['text'])
                doc.add_page_break()
                
            doc.save(self.output_path)
            logger.info(f"DOCX created: {self.output_path}")
            return True
            
        except ImportError:
            logger.error("python-docx is not installed. Install it with: pip install python-docx")
            return False
        except Exception as e:
            logger.error(f"Error creating DOCX: {e}")
            return False
    
    def run(self):
        """Main execution flow of the converter"""
        logger.info(f"Starting conversion of {self.base_url} to {self.options['format']}")
        
        # Use sitemap if requested
        if self.options['use_sitemap']:
            logger.info("Using sitemap for URL discovery")
            urls = self.parse_sitemap(self.options['sitemap_url'])
            self.to_visit = urls
        
        # Interactive mode
        if self.options['interactive']:
            logger.info("Running in interactive mode")
            self.to_visit = self.interactive_mode()
        
        # Main crawling and processing loop
        while self.to_visit and len(self.visited_urls) < self.options['max_pages']:
            current_url = self.to_visit.pop(0)
            
            if current_url in self.visited_urls:
                continue
                
            logger.info(f"Processing {len(self.visited_urls) + 1}/{self.options['max_pages']}: {current_url}")
            self.visited_urls.add(current_url)
            
            elements = self.process_page(current_url)
            if elements:
                self.pdf_elements.extend(elements)
                
            # Respect the delay between requests
            time.sleep(self.options['delay'])
        
        # Create the requested output format
        if self.options['format'].lower() == 'pdf':
            return self.create_pdf()
        elif self.options['format'].lower() == 'html':
            return self.create_html()
        elif self.options['format'].lower() == 'json':
            return self.create_json()
        elif self.options['format'].lower() == 'md' or self.options['format'].lower() == 'markdown':
            return self.create_markdown()
        elif self.options['format'].lower() == 'docx':
            return self.create_docx()
        else:
            logger.error(f"Unsupported format: {self.options['format']}")
            return False


def main():
    """Parse command line arguments and run the converter"""
    parser = argparse.ArgumentParser(description="Web to PDF Converter Tool")
    
    parser.add_argument("--url", required=True, help="URL to start crawling from")
    parser.add_argument("--output", help="Output file path (default: output.[format])")
    parser.add_argument("--format", choices=["pdf", "html", "json", "docx", "md", "markdown"], default="pdf", 
                        help="Output format (default: pdf)")
    parser.add_argument("--max-depth", type=int, help="Maximum depth to crawl (default: no limit)")
    parser.add_argument("--max-pages", type=int, default=250, help="Maximum number of pages to process (default: 250)")
    parser.add_argument("--delay", type=float, default=1, help="Delay between requests in seconds (default: 1)")
    parser.add_argument("--timeout", type=int, default=10, help="Request timeout in seconds (default: 10)")
    parser.add_argument("--contains", help="Only include pages containing these keywords (comma-separated)")
    parser.add_argument("--not-contains", help="Exclude pages containing these keywords (comma-separated)")
    parser.add_argument("--categories", help="Only include pages from these categories (comma-separated)")
    parser.add_argument("--use-sitemap", action="store_true", help="Use sitemap for URL discovery")
    parser.add_argument("--sitemap-url", help="URL of the sitemap (default: auto-detect)")
    parser.add_argument("--toc", action="store_true", help="Generate table of contents")
    parser.add_argument("--interactive", action="store_true", help="Interactive mode for URL selection")
    
    args = parser.parse_args()
    
    # Determine output file path
    if not args.output:
        args.output = f"output.{args.format}"
    elif not args.output.endswith(f".{args.format}"):
        args.output = f"{args.output}.{args.format}"
    
    # Create converter options
    options = {
        'max_pages': args.max_pages,
        'max_depth': args.max_depth,
        'delay': args.delay,
        'timeout': args.timeout,
        'format': args.format,
        'use_sitemap': args.use_sitemap,
        'sitemap_url': args.sitemap_url,
        'contains': args.contains,
        'not_contains': args.not_contains,
        'categories': args.categories,
        'create_toc': args.toc,
        'interactive': args.interactive
    }
    
    # Create and run the converter
    converter = WebToPdfConverter(args.url, args.output, options)
    success = converter.run()
    
    if success:
        logger.info("Conversion completed successfully.")
    else:
        logger.error("Conversion failed.")


if __name__ == "__main__":
    main()